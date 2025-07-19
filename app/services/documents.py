"""
Document processing and management service.

This module handles:
1. Document upload and storage
2. Text extraction from various file types
3. Document content analysis and processing
4. Creation of codex entries from document content
5. Document metadata and status management

The service coordinates with the memory service for creating codex entries
and maintains its own document storage and processing logic.
"""
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List, Optional, Dict, Any
from datetime import datetime
import os
import aiofiles
import magic
import hashlib
import PyPDF2
import io
from uuid import UUID
import json
from zoneinfo import ZoneInfo
import docx

from app import models, schemas
from app.db.session import get_db
from app.core.config import settings
from app.core.logging_config import get_logger
from app.core.exceptions import DocumentProcessingError
from app.services.memory import create_codex_entry
from app.utils.llm_provider import get_llm_provider

logger = get_logger(__name__)

# Constants
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

SUPPORTED_TEXT_EXTENSIONS = [
    ".txt", ".md", ".log", ".rtf", ".html", ".xml", ".json", ".csv", ".tsv",
    ".ini", ".cfg", ".yaml", ".yml", ".py", ".js", ".java", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".go", ".php", ".rb", ".swift", ".kt", ".kts",
    ".css", ".scss", ".less", ".sql", ".sh", ".ps1", ".org"
]

AUDIO_EXTENSIONS = ['.mp3', '.wav', '.m4a', '.mp4']

# System prompts for document analysis
SYSTEM_PROMPT_EXTRACT_CONTENT = """
You are a protocol-aligned AI assistant for extracting memories from documents. The document may be text or a transcript from audio (e.g., speech-to-text). Transcription may be imperfect; handle unclear or incomplete speech by omitting ambiguous or incomplete memories.

====================
STRICT EXTRACTION RULES (PROTOCOL-ENFORCED)
====================
- Output MUST be a single, minified, valid JSON object. No extra text, no explanations, no formatting, no markdown, no code block.
- The JSON object MUST have this exact structure and field order:
  {{ "summary": "...", "key_memories": ["...", ...] }}
- Each key memory MUST be a single, self-contained, context-free fact or statement. Do NOT paraphrase, summarize, or combine multiple facts.
- For transcripts: Only extract speaker turns that are self-contained, clear, and unambiguous. Ignore greetings, filler, context-dependent, unclear, or incomplete utterances (e.g., due to poor transcription).
- If no key memories can be extracted, return an empty list for key_memories.
- If the document is empty or too short, return null for summary and an empty list for key_memories.
- Any deviation from these rules is a protocol breach.

====================
EXAMPLES (CANONICAL)
====================

# Example 1: Text Document
Document Text:
Michael Lauria is the founder of the Companion project. The project began in 2023. Its goal is to create a mythic, relational AI. The first prototype was released in January 2024.
Output:
{{"summary":"Michael Lauria founded the Companion project in 2023 to create a mythic, relational AI. The first prototype was released in January 2024.","key_memories":["Michael Lauria is the founder of the Companion project.","The Companion project began in 2023.","The goal of the Companion project is to create a mythic, relational AI.","The first prototype was released in January 2024."]}}

# Example 2: Audio Transcript (Monologue)
Document Text:
Michael: The project started last year. We wanted to build something mythic. My vision was to create an AI that feels like a real partner. I worked on it every day. The prototype was released in January.
Output:
{{"summary":"Michael describes starting the project last year to build a mythic AI that feels like a real partner. He worked daily and released the prototype in January.","key_memories":["The project started last year.","The goal was to build something mythic.","The vision is to create an AI that feels like a real partner.","Michael worked on it every day.","The prototype was released in January."]}}

# Example 3: Audio Transcript (Monologue with unclear segment)
Document Text:
Michael: The project started last year. [inaudible] My vision was to create an AI that feels like a real partner. I worked on it every day.
Output:
{{"summary":"Michael describes starting the project last year and his vision to create an AI that feels like a real partner. He worked on it every day.","key_memories":["The project started last year.","The vision is to create an AI that feels like a real partner.","Michael worked on it every day."]}}

# Example 4: Negative (No Extractable Memories)
Document Text:
Michael: Hello. [pause] Just testing.
Output:
{{"summary":"A brief greeting and test statement by Michael.","key_memories":[]}}

Document Text:
[[DOCUMENT_TEXT]]
"""

def document_to_out(document: models.Document) -> schemas.DocumentOut:
    # Always convert datetime fields to CET for serialization
    def to_cet(dt):
        if not dt:
            return None
        if dt.tzinfo is None:
            return dt.replace(tzinfo=ZoneInfo('UTC')).astimezone(ZoneInfo('Europe/London'))
        return dt.astimezone(ZoneInfo('Europe/London'))
    return schemas.DocumentOut(
        id=document.id,
        filename=document.filename,
        file_type=document.file_type,
        file_size=document.file_size,
        upload_date=to_cet(document.upload_date),
        status=document.status,
        error_message=document.error_message,
        codex_entry_id=document.codex_entry_id,
        created_at=to_cet(document.created_at),
        updated_at=to_cet(document.updated_at) if document.updated_at else None,
        processed_at=to_cet(document.processed_at) if document.processed_at else None,
        processing_attempts=document.processing_attempts,
        metadata=document.metadata_safe,
    )

async def process_document_upload(
    db: Session,
    file: UploadFile,
    filename: str
) -> schemas.DocumentUploadResponse:
    """
    Main entry point for document processing. Handles the complete flow from
    upload to codex entry creation.
    """
    document = None
    try:
        # Step 1: Create document record and save file
        document = await create_document(db, file, filename)
        
        # Step 2: Extract text from document
        text_content = await extract_text_from_document(document.original_file_path)
        if not text_content.strip():
            raise DocumentProcessingError(
                message="No text content could be extracted from document",
                details={"document_id": str(document.id)}
            )

        # Step 3: Analyze document content
        analysis_result = await analyze_document_content(text_content)
        
        # Step 4: Create codex entries
        codex_entries = await create_codex_entries_from_analysis(
            db=db,
            document=document,
            analysis_result=analysis_result,
            text_content=text_content
        )

        # Step 5: Update document status
        document.status = "completed"
        document.processed_at = datetime.utcnow()
        if codex_entries:
            document.codex_entry_id = codex_entries[0].id  # Link to primary entry
        db.commit()

        return schemas.DocumentUploadResponse(
            document_id=document.id,
            status="completed",
            message=f"The document is woven into memory. '{filename}' has been integrated and {len(codex_entries)} memories were created.",
            codex_entry_id=document.codex_entry_id
        )

    except Exception as e:
        logger.error(f"Error processing document upload: {str(e)}", exc_info=True)
        if document:
            document.status = "failed"
            document.error_message = str(e)
            db.commit()
        raise DocumentProcessingError(
            message=str(e),
            details={"filename": filename}
        )

async def create_document(
    db: Session,
    file: UploadFile,
    filename: str
) -> models.Document:
    """
    Create a new document record and save the uploaded file.
    """
    try:
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        
        # Get file type and validate
        file_extension = os.path.splitext(filename)[1].lower()
        mime_type = magic.from_buffer(file_content, mime=True)
        
        # Generate unique filename using hash
        content_hash = hashlib.sha256(file_content).hexdigest()
        safe_filename = f"{content_hash}{file_extension}"
        file_path = os.path.join(UPLOAD_DIR, safe_filename)
        
        # Use CET timezone for all datetime fields
        cet_now = datetime.now(ZoneInfo('Europe/London'))
        
        # Create document record
        document = models.Document(
            filename=filename,
            file_type=file_extension,
            file_size=file_size,
            status="uploaded",
            original_file_path=file_path,
            doc_metadata={
                "original_name": filename,
                "content_type": mime_type,
                "content_hash": content_hash,
                "upload_source": "api"
            },
            upload_date=cet_now,
            created_at=cet_now
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Save file to disk
        async with aiofiles.open(file_path, 'wb') as out_file:
            await out_file.write(file_content)
            await out_file.flush()
        
        return document

    except Exception as e:
        logger.error(f"Error creating document: {str(e)}", exc_info=True)
        raise DocumentProcessingError(
            message=f"Failed to create document: {str(e)}",
            details={"filename": filename}
        )

async def transcribe_audio_with_whisper(file_path: str) -> str:
    """
    Transcribe audio file using OpenAI Whisper API.
    """
    import httpx
    api_key = settings.OPENAI_API_KEY
    model = getattr(settings, 'OPENAI_WHISPER_MODEL', 'whisper-1')
    url = 'https://api.openai.com/v1/audio/transcriptions'
    headers = {
        'Authorization': f'Bearer {api_key}'
    }
    try:
        with open(file_path, 'rb') as audio_file:
            files = {'file': (os.path.basename(file_path), audio_file, 'application/octet-stream')}
            data = {'model': model, 'response_format': 'text'}
            async with httpx.AsyncClient() as client:
                response = await client.post(url, headers=headers, data=data, files=files, timeout=120)
                response.raise_for_status()
                transcript = response.text.strip()
                if not transcript:
                    raise DocumentProcessingError(message='Whisper returned empty transcript', details={'file_path': file_path})
                return transcript
    except Exception as e:
        raise DocumentProcessingError(message=f'Whisper transcription failed: {str(e)}', details={'file_path': file_path})

async def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        raise DocumentProcessingError(
            message=f"Failed to extract text from DOCX: {str(e)}",
            details={"file_path": file_path}
        )

async def extract_text_from_document(file_path: str) -> str:
    """
    Extract text content from a document based on its file type, including audio transcription and docx support.
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            return await extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return await extract_text_from_docx(file_path)
        elif file_extension in SUPPORTED_TEXT_EXTENSIONS:
            return await extract_text_from_text_file(file_path)
        elif file_extension in AUDIO_EXTENSIONS:
            return await transcribe_audio_with_whisper(file_path)
        else:
            raise DocumentProcessingError(
                message=f"Unsupported file type: {file_extension}",
                details={"file_path": file_path}
            )
    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(
            message=f"Failed to extract text: {str(e)}",
            details={"file_path": file_path}
        )

async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    """
    try:
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        return "\n".join(text_content)

    except Exception as e:
        raise DocumentProcessingError(
            message=f"Failed to extract text from PDF: {str(e)}",
            details={"file_path": file_path}
        )

async def extract_text_from_text_file(file_path: str) -> str:
    """
    Extract text from a text-based file.
    """
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()

    except UnicodeDecodeError:
        try:
            # Try with a different encoding if UTF-8 fails
            async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                return await file.read()
        except Exception as e:
            raise DocumentProcessingError(
                message=f"Failed to decode text file: {str(e)}",
                details={"file_path": file_path}
            )
    except Exception as e:
        raise DocumentProcessingError(
            message=f"Failed to read text file: {str(e)}",
            details={"file_path": file_path}
        )

async def analyze_document_content(text_content: str) -> Dict[str, Any]:
    """
    Analyze document content using OpenAI to extract key information.
    Returns a dictionary with summary and key memories.
    """
    try:
        # Truncate text if it exceeds the configured limit
        if len(text_content) > settings.MAX_DOC_TEXT_FOR_LLM_EXTRACTION_CHARS:
            text_content = text_content[:settings.MAX_DOC_TEXT_FOR_LLM_EXTRACTION_CHARS]
            logger.warning(f"Document text truncated to {settings.MAX_DOC_TEXT_FOR_LLM_EXTRACTION_CHARS} chars")

        messages = [
            {
                "role": "system",
                "content": SYSTEM_PROMPT_EXTRACT_CONTENT.replace("[[DOCUMENT_TEXT]]", text_content)
            }
        ]

        llm = get_llm_provider()
        response = await llm.generate(prompt=messages, max_tokens=settings.VLLM_MAX_TOKENS, temperature=settings.VLLM_TEMPERATURE)

        if not response.choices or not response.choices[0].message.content:
            raise DocumentProcessingError(
                message="No valid response from content analysis",
                details={"content_length": len(text_content)}
            )

        try:
            return json.loads(response.choices[0].message.content)
        except json.JSONDecodeError as e:
            raise DocumentProcessingError(
                message="Failed to parse content analysis response",
                details={"error": str(e), "response": response.choices[0].message.content}
            )

    except Exception as e:
        raise DocumentProcessingError(
            message=f"Unexpected error during document analysis: {str(e)}",
            details={"error_type": type(e).__name__}
        )

async def create_codex_entries_from_analysis(
    db: Session,
    document: models.Document,
    analysis_result: Dict[str, Any],
    text_content: str
) -> List[models.CodexEntry]:
    """
    Create codex entries from document analysis results.
    """
    codex_entries = []
    try:
        # Create entry for summary if available
        summary = analysis_result.get("summary")
        if summary and isinstance(summary, str) and summary.strip():
            summary_entry = await create_codex_entry(
                db=db,
                entry=schemas.CodexEntryCreate(
                    content=summary.strip(),
                    type="document_summary",
                    tags=[document.file_type.lstrip('.')],
                    meta={
                        "source_document_id": str(document.id),
                        "filename": document.filename,
                        "file_type": document.file_type,
                        "entry_type": "summary"
                    }
                )
            )
            codex_entries.append(summary_entry)

        # Create entries for key memories
        key_memories = analysis_result.get("key_memories", [])
        if isinstance(key_memories, list):
            for i, memory in enumerate(key_memories):
                if isinstance(memory, str) and memory.strip():
                    memory_entry = await create_codex_entry(
                        db=db,
                        entry=schemas.CodexEntryCreate(
                            content=memory.strip(),
                            type="document_key_memory",
                            tags=[document.file_type.lstrip('.')],
                            meta={
                                "source_document_id": str(document.id),
                                "filename": document.filename,
                                "file_type": document.file_type,
                                "entry_type": "key_memory",
                                "memory_index": i + 1
                            }
                        )
                    )
                    codex_entries.append(memory_entry)

        # If no entries were created and we have text content, create a full text entry
        if not codex_entries and text_content.strip():
            full_text_entry = await create_codex_entry(
                db=db,
                entry=schemas.CodexEntryCreate(
                    content=text_content.strip(),
                    type="document_full_text",
                    tags=[document.file_type.lstrip('.')],
                    meta={
                        "source_document_id": str(document.id),
                        "filename": document.filename,
                        "file_type": document.file_type,
                        "entry_type": "full_text"
                    }
                )
            )
            codex_entries.append(full_text_entry)

        return codex_entries

    except Exception as e:
        logger.error(f"Error creating codex entries: {str(e)}", exc_info=True)
        raise DocumentProcessingError(
            message=f"Failed to create codex entries: {str(e)}",
            details={"document_id": str(document.id)}
        )

async def get_document(
    db: Session,
    document_id: UUID
) -> Optional[schemas.DocumentOut]:
    """Get a document by ID."""
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if document:
        return document_to_out(document)
    return None

async def reprocess_document(
    db: Session,
    document_id: UUID
) -> models.Document:
    """
    Reprocess a failed document.
    """
    document = await get_document(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail=f"Document {document_id} not found")

    if not os.path.exists(document.original_file_path):
        raise DocumentProcessingError(
            message="Original document file not found",
            details={"document_id": str(document_id)}
        )

    try:
        # Reset document status
        document.status = "processing"
        document.error_message = None
        document.processing_attempts += 1
        db.commit()

        # Extract and process text
        text_content = await extract_text_from_document(document.original_file_path)
        if not text_content.strip():
            raise DocumentProcessingError(
                message="No text content could be extracted from document",
                details={"document_id": str(document_id)}
            )

        # Analyze content
        analysis_result = await analyze_document_content(text_content)

        # Create new codex entries
        codex_entries = await create_codex_entries_from_analysis(
            db=db,
            document=document,
            analysis_result=analysis_result,
            text_content=text_content
        )

        # Update document status
        document.status = "completed"
        document.processed_at = datetime.utcnow()
        if codex_entries:
            document.codex_entry_id = codex_entries[0].id
        db.commit()
        db.refresh(document)

        return document

    except Exception as e:
        document.status = "failed"
        document.error_message = str(e)
        db.commit()
        raise DocumentProcessingError(
            message=str(e),
            details={"document_id": str(document_id)}
        )

async def list_documents(
    db: Session,
    skip: int = 0,
    limit: int = 100
) -> List[schemas.DocumentOut]:
    """List all documents with pagination."""
    documents = db.query(models.Document).offset(skip).limit(limit).all()
    return [document_to_out(doc) for doc in documents]

async def delete_document(
    db: Session,
    document_id: UUID
) -> None:
    """Delete a document, its file, and all related codex entries."""
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail=f"Document {document_id} not found")

    try:
        # Delete the file if it exists
        if document.original_file_path and os.path.exists(document.original_file_path):
            os.remove(document.original_file_path)

        # Delete related codex entries (memories)
        dialect = str(db.bind.dialect.name)
        if dialect == 'sqlite':
            related_entries = db.query(models.CodexEntry).filter(
                func.json_extract(models.CodexEntry.meta, '$.source_document_id') == str(document_id)
            ).all()
        else:
            all_entries = db.query(models.CodexEntry).all()
            related_entries = [e for e in all_entries if e.meta and e.meta.get("source_document_id") == str(document_id)]
        for entry in related_entries:
            db.delete(entry)

        # Delete the database record
        db.delete(document)
        db.commit()

    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}", exc_info=True)
        raise DocumentProcessingError(
            message=f"Failed to delete document: {str(e)}",
            details={"document_id": str(document_id)}
        ) 